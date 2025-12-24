"""
Modulo Generazione Relazione Tecnica Avanzata
Genera relazioni tecniche complete con formule, diagrammi e disegni
Supporto completo per aperture ad arco
Arch. Michelangelo Bartolotta
"""

from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from datetime import datetime
import os
from typing import Dict
import json
import math

# Per generazione Word avanzata
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx non installato")

# Per grafici
try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib.patches import Rectangle, FancyBboxPatch, Arc
    import matplotlib.patches as mpatches
    from io import BytesIO
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False
    print("matplotlib non installato")


class EnhancedReportGenerator:
    """Generatore avanzato di relazioni tecniche con supporto aperture ad arco"""
    
    def __init__(self):
        self.project_data = {}
        self.results = {}
        self.engineer_data = {
            'nome': 'Arch. Michelangelo Bartolotta',
            'titolo': 'Architetto',
            'iscrizione': 'Ordine degli Architetti di Agrigento n. 1557',
            'cf': 'BRTMHL68C19A089V',
            'piva': '02554080842',
            'sede_op': 'Via Domenico Scinà n. 28, Palermo',
            'sede_legale': 'C.le S.to Argento n.11, 92020 Castrofilippo (AG)',
            'email': 'arch.bartolotta@pec.it',
            'tel': '+39 XXX XXXXXXX'
        }
        
    def set_data(self, project_data: Dict, results: Dict):
        """Imposta i dati del progetto e i risultati"""
        self.project_data = project_data
        self.results = results
        
    def generate_enhanced_word_report(self, filename: str):
        """Genera relazione Word completa con formule, diagrammi e disegni"""
        if not DOCX_AVAILABLE:
            raise Exception("Modulo python-docx non disponibile")
            
        doc = Document()
        
        # Imposta margini
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        
        # Intestazione professionale
        self._add_professional_header(doc)
        
        # Frontespizio
        self._add_cover_page(doc)
        
        # Indice
        doc.add_page_break()
        self._add_table_of_contents(doc)
        
        # 1. Premessa
        doc.add_page_break()
        self._add_introduction(doc)
        
        # 2. Riferimenti normativi
        self._add_normative_references(doc)
        
        # 3. Descrizione dell'intervento
        self._add_intervention_description(doc)
        
        # 4. Caratteristiche dei materiali
        self._add_materials_section(doc)
        
        # 5. Geometria della struttura
        self._add_geometry_section(doc)
        
        # 6. Analisi dei carichi
        self._add_loads_analysis(doc)
        
        # 7. Procedura di calcolo
        self._add_calculation_procedure(doc)
        
        # 8. Verifiche stato di fatto
        self._add_existing_state_verification(doc)
        
        # 9. Verifiche stato di progetto
        self._add_project_state_verification(doc)
        
        # 10. Confronto e verifiche normative
        self._add_normative_verifications(doc)
        
        # 11. Conclusioni
        self._add_conclusions(doc)
        
        # Allegati
        doc.add_page_break()
        self._add_attachments(doc)
        
        # Salva documento
        doc.save(filename)
        
    def _add_professional_header(self, doc):
        """Aggiunge intestazione professionale completa"""
        # Crea tabella per intestazione
        table = doc.add_table(rows=1, cols=3)
        table.style = None
        
        # Logo/Timbro professionale (placeholder)
        cell_left = table.rows[0].cells[0]
        p = cell_left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("STUDIO TECNICO\n").bold = True
        p.add_run("di ARCHITETTURA\n").italic = True
        
        # Dati centrali
        cell_center = table.rows[0].cells[1]
        p = cell_center.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(self.engineer_data['nome']).bold = True
        p.add_run(f"\n{self.engineer_data['titolo']}")
        p.add_run(f"\n{self.engineer_data['iscrizione']}")
        
        # Contatti
        cell_right = table.rows[0].cells[2]
        p = cell_right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"{self.engineer_data['sede_op']}")
        p.add_run(f"\nC.F. {self.engineer_data['cf']}")
        p.add_run(f"\nP.IVA {self.engineer_data['piva']}")
        
        # Linea separatrice
        doc.add_paragraph("_" * 80)
        
    def _add_cover_page(self, doc):
        """Aggiunge frontespizio"""
        # Spazio iniziale
        for _ in range(5):
            doc.add_paragraph()
            
        # Titolo principale
        title = doc.add_heading('RELAZIONE TECNICA DI CALCOLO', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        subtitle = doc.add_heading('VERIFICA LOCALE PER APERTURA IN PARETE MURARIA', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle2 = doc.add_heading('ai sensi del D.M. 17/01/2018 § 8.4.1', 2)
        subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Spazio
        for _ in range(3):
            doc.add_paragraph()
            
        # Info progetto
        info = self.project_data.get('info', {})
        
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = None
        
        data = [
            ("OGGETTO:", info.get('name', 'Intervento di cerchiatura')),
            ("COMMITTENTE:", info.get('client', 'Da definire')),
            ("UBICAZIONE:", info.get('location', 'Da definire')),
            ("DATA:", info.get('date', datetime.now().strftime('%d/%m/%Y')))
        ]
        
        for i, (label, value) in enumerate(data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            info_table.rows[i].cells[1].text = value
            
    def _add_table_of_contents(self, doc):
        """Aggiunge indice"""
        doc.add_heading('INDICE', 1)
        
        toc_items = [
            ("1.", "PREMESSA", "3"),
            ("2.", "RIFERIMENTI NORMATIVI", "4"),
            ("3.", "DESCRIZIONE DELL'INTERVENTO", "5"),
            ("4.", "CARATTERISTICHE DEI MATERIALI", "6"),
            ("5.", "GEOMETRIA DELLA STRUTTURA", "8"),
            ("6.", "ANALISI DEI CARICHI", "10"),
            ("7.", "PROCEDURA DI CALCOLO", "12"),
            ("8.", "VERIFICHE STATO DI FATTO", "15"),
            ("9.", "VERIFICHE STATO DI PROGETTO", "18"),
            ("10.", "CONFRONTO E VERIFICHE NORMATIVE", "21"),
            ("11.", "CONCLUSIONI", "23"),
            ("", "ALLEGATI", "24"),
            ("A.", "Disegni tecnici", "24"),
            ("B.", "Diagrammi di calcolo", "26"),
            ("C.", "Tabulati di calcolo", "28")
        ]
        
        for num, title, page in toc_items:
            p = doc.add_paragraph()
            p.add_run(f"{num} {title}").bold = (num != "")
            p.add_run(f" {'.' * (60 - len(num) - len(title))} {page}")
            
    def _add_introduction(self, doc):
        """Aggiunge premessa"""
        doc.add_heading('1. PREMESSA', 1)
        
        p = doc.add_paragraph()
        p.add_run("La presente relazione tecnica riguarda la verifica strutturale locale per la ")
        p.add_run("realizzazione di aperture in parete muraria portante").bold = True
        p.add_run(", con inserimento di cerchiature metalliche di rinforzo.")
        
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run("L'intervento è inquadrato nell'ambito degli ")
        p.add_run("interventi locali").bold = True
        p.add_run(" ai sensi del D.M. 17/01/2018 § 8.4.1, in quanto finalizzato a:")
        
        items = [
            "Ripristinare, rispetto alla configurazione precedente, le caratteristiche iniziali di elementi danneggiati",
            "Migliorare le caratteristiche di resistenza e/o duttilità di elementi strutturali",
            "Impedire meccanismi di collasso locale",
            "Modificare un elemento senza alterare significativamente il comportamento globale"
        ]
        
        for item in items:
            doc.add_paragraph(f"• {item}", style='List Bullet')
            
    def _add_normative_references(self, doc):
        """Aggiunge riferimenti normativi dettagliati"""
        doc.add_heading('2. RIFERIMENTI NORMATIVI', 1)
        
        doc.add_heading('2.1 Normativa nazionale', 2)
        
        norms = [
            ("D.M. 17/01/2018", "Aggiornamento delle «Norme tecniche per le costruzioni»"),
            ("Circolare 21/01/2019 n. 7", "Istruzioni per l'applicazione dell'«Aggiornamento delle 'Norme tecniche per le costruzioni'» di cui al D.M. 17 gennaio 2018"),
            ("D.P.R. 380/2001", "Testo unico delle disposizioni legislative e regolamentari in materia edilizia")
        ]
        
        for code, desc in norms:
            p = doc.add_paragraph()
            p.add_run(f"• {code}").bold = True
            p.add_run(f" - {desc}")
            
        doc.add_heading('2.2 Riferimenti specifici NTC 2018', 2)
        
        refs = [
            "§ 8.4.1 - Riparazione o intervento locale",
            "§ 8.7.1 - Costruzioni in muratura",
            "§ 7.8.1.5.2 - Verifiche alle tensioni ammissibili",
            "§ 4.5.6.3 - Verifiche agli stati limite ultimi"
        ]
        
        for ref in refs:
            doc.add_paragraph(f"• {ref}", style='List Bullet')
            
    def _add_intervention_description(self, doc):
        """Aggiunge descrizione dettagliata dell'intervento con supporto aperture ad arco"""
        doc.add_heading('3. DESCRIZIONE DELL\'INTERVENTO', 1)
        
        doc.add_heading('3.1 Descrizione generale', 2)
        
        openings = self.project_data.get('openings_module', {}).get('openings', [])
        n_openings = len(openings)
        
        # Conta aperture per tipo
        n_rect = sum(1 for op in openings if op.get('type', 'Rettangolare') == 'Rettangolare')
        n_arch = sum(1 for op in openings if op.get('type', 'Rettangolare') == 'Ad arco')
        
        p = doc.add_paragraph()
        p.add_run("L'intervento prevede la realizzazione di ")
        p.add_run(f"n. {n_openings} apertur{'a' if n_openings == 1 else 'e'}").bold = True
        
        # Specifica il tipo di aperture
        if n_arch > 0 and n_rect == 0:
            p.add_run(" AD ARCO")
        elif n_rect > 0 and n_arch == 0:
            p.add_run(" RETTANGOLARI")
        elif n_arch > 0 and n_rect > 0:
            p.add_run(f" ({n_rect} rettangolar{'e' if n_rect == 1 else 'i'} e {n_arch} ad arco)")
        
        p.add_run(" in parete muraria portante esistente, mediante:")
        
        steps = [
            "Puntellamento provvisionale della parete",
            "Demolizione controllata della muratura",
            "Inserimento di cerchiature metalliche (architrave e piedritti)",
            "Ripristino delle connessioni murarie",
            "Finitura e protezione degli elementi metallici"
        ]
        
        # Se ci sono aperture ad arco, aggiungi step specifico
        if n_arch > 0:
            steps.insert(3, "Calandratura dei profili metallici per aperture ad arco (ove necessario)")
        
        for step in steps:
            doc.add_paragraph(f"• {step}", style='List Bullet')
            
        # Nota specifica per aperture ad arco
        if n_arch > 0:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Nota: ").bold = True
            p.add_run("Le aperture ad arco richiedono particolare attenzione nella realizzazione delle cerchiature, ")
            p.add_run("che devono seguire la geometria curvilinea dell'apertura. I profili metallici devono essere ")
            p.add_run("opportunamente calandrati o sostituiti con elementi idonei alla curvatura richiesta.")
            
        doc.add_heading('3.2 Classificazione dell\'intervento', 2)
        
        p = doc.add_paragraph()
        p.add_run("Ai sensi del § 8.4.1 delle NTC 2018, l'intervento è classificabile come ")
        p.add_run("INTERVENTO LOCALE").bold = True
        p.add_run(" quando soddisfa i seguenti requisiti:")
        
        # Inserisci formula box
        self._add_formula_box(doc, 
            "Variazione di rigidezza: |ΔK| ≤ 15%\n" +
            "Variazione di resistenza: ΔV ≤ 20%"
        )
        
    def _add_materials_section(self, doc):
        """Aggiunge sezione materiali con tabelle dettagliate"""
        doc.add_heading('4. CARATTERISTICHE DEI MATERIALI', 1)
        
        doc.add_heading('4.1 Muratura esistente', 2)
        
        masonry = self.project_data.get('masonry', {})
        
        # Tabella proprietà muratura
        table = doc.add_table(rows=8, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Intestazioni
        headers = ['Parametro', 'Simbolo', 'Valore']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            
        # Dati muratura
        data = [
            ('Tipologia', '-', masonry.get('type', 'Muratura in pietrame')),
            ('Resistenza media a compressione', 'f_cm', f"{masonry.get('fcm', 3.2)} MPa"),
            ('Resistenza media a taglio', 'τ_0', f"{masonry.get('tau0', 0.076)} MPa"),
            ('Modulo elastico normale', 'E', f"{masonry.get('E', 1740)} MPa"),
            ('Modulo elastico tangenziale', 'G', f"{masonry.get('E', 1740)/6:.0f} MPa"),
            ('Peso specifico', 'γ', f"{masonry.get('gamma', 19)} kN/m³"),
            ('Livello di conoscenza', 'LC', masonry.get('knowledge_level', 'LC1'))
        ]
        
        for i, (param, symbol, value) in enumerate(data, start=1):
            table.rows[i].cells[0].text = param
            table.rows[i].cells[1].text = symbol
            table.rows[i].cells[2].text = value
            
        doc.add_heading('4.2 Acciaio per cerchiature', 2)
        
        p = doc.add_paragraph()
        p.add_run("Per le cerchiature metalliche si prevede l'utilizzo di:")
        
        steel_props = [
            "Profili HEA/HEB/IPE in acciaio S275",
            "Tensione di snervamento: f_yk = 275 MPa",
            "Tensione di rottura: f_tk = 430 MPa",
            "Modulo elastico: E_s = 210000 MPa",
            "Coefficiente di sicurezza: γ_M0 = 1.05"
        ]
        
        for prop in steel_props:
            doc.add_paragraph(f"• {prop}", style='List Bullet')
            
        # Aggiungi nota per aperture ad arco
        openings = self.project_data.get('openings_module', {}).get('openings', [])
        n_arch = sum(1 for op in openings if op.get('type', 'Rettangolare') == 'Ad arco')
        
        if n_arch > 0:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Nota per aperture ad arco: ").bold = True
            p.add_run("I profili destinati alle aperture ad arco devono essere sottoposti a calandratura ")
            p.add_run("a freddo per adeguarsi alla geometria curvilinea. In caso di profili non calandrabili, ")
            p.add_run("si prevede l'utilizzo di soluzioni alternative (profili segmentati, piastre sagomate, etc.).")
            
    def _add_geometry_section(self, doc):
        """Aggiunge sezione geometria con disegni e supporto aperture ad arco"""
        doc.add_heading('5. GEOMETRIA DELLA STRUTTURA', 1)
        
        doc.add_heading('5.1 Geometria della parete', 2)
        
        wall = self.project_data.get('wall', {})
        
        # Tabella dimensioni
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Light Grid Accent 1'
        
        headers = ['Dimensione', 'Simbolo', 'Valore']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            
        data = [
            ('Lunghezza totale', 'L', f"{wall.get('length', 500)} cm"),
            ('Altezza', 'h', f"{wall.get('height', 300)} cm"),
            ('Spessore', 't', f"{wall.get('thickness', 50)} cm")
        ]
        
        for i, (dim, symbol, value) in enumerate(data, start=1):
            table.rows[i].cells[0].text = dim
            table.rows[i].cells[1].text = symbol
            table.rows[i].cells[2].text = value
            
        doc.add_heading('5.2 Aperture previste', 2)
        
        openings = self.project_data.get('openings_module', {}).get('openings', [])
        
        if openings:
            # Tabella con colonna tipo
            table = doc.add_table(rows=len(openings)+1, cols=6)
            table.style = 'Light Grid Accent 1'
            
            headers = ['Apertura', 'Tipo', 'Larghezza [cm]', 'Altezza [cm]', 'Posizione X [cm]', 'Quota base [cm]']
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
                
            for i, opening in enumerate(openings):
                table.rows[i+1].cells[0].text = f"AP{i+1}"
                table.rows[i+1].cells[1].text = opening.get('type', 'Rettangolare')
                table.rows[i+1].cells[2].text = str(opening.get('width', 120))
                
                # Gestione altezza per aperture ad arco
                height_text = str(opening.get('height', 240))
                if opening.get('type') == 'Ad arco' and 'arch_params' in opening:
                    params = opening['arch_params']
                    rise = params.get('rise', 0)
                    if rise > 0:
                        height_text += f"\n(freccia: {rise} cm)"
                
                table.rows[i+1].cells[3].text = height_text
                table.rows[i+1].cells[4].text = str(opening.get('x', 0))
                table.rows[i+1].cells[5].text = str(opening.get('base_height', 0))
                
        # Dettagli cerchiature per aperture
        doc.add_heading('5.3 Dettagli cerchiature', 2)
        
        for i, opening in enumerate(openings):
            if 'reinforcement' in opening and opening['reinforcement']:
                p = doc.add_paragraph()
                p.add_run(f"Apertura AP{i+1}: ").bold = True
                
                reinf = opening['reinforcement']
                op_type = opening.get('type', 'Rettangolare')
                
                # Tipo cerchiatura
                p.add_run(f"{reinf.get('type', 'N/A')}")
                
                # Materiale e profili
                if reinf.get('material') == 'acciaio' and 'architrave' in reinf:
                    arch = reinf['architrave']
                    profilo = arch.get('profilo', 'N/A')
                    p.add_run(f" - Profilo: {profilo}")
                    
                    # Specificità per aperture ad arco
                    if op_type == 'Ad arco':
                        if arch.get('curvato', False):
                            p.add_run(" (calandrato)")
                        else:
                            p.add_run(" (non calandrabile - soluzione alternativa richiesta)")
                
        # Disegni tecnici
        if MATPLOTLIB_AVAILABLE:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Figura 5.1 - Schema geometrico stato di fatto").italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Genera e inserisci disegno stato di fatto
            img_path = self._generate_wall_drawing(state='existing')
            if img_path and os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(5))
                
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Figura 5.2 - Schema geometrico stato di progetto").italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Genera e inserisci disegno stato di progetto
            img_path = self._generate_wall_drawing(state='project')
            if img_path and os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(5))
                
    def _add_loads_analysis(self, doc):
        """Aggiunge analisi dei carichi dettagliata"""
        doc.add_heading('6. ANALISI DEI CARICHI', 1)
        
        doc.add_heading('6.1 Carichi verticali', 2)
        
        loads = self.project_data.get('loads', {})
        
        p = doc.add_paragraph()
        p.add_run("I carichi verticali agenti sulla parete sono stati determinati considerando:")
        
        doc.add_paragraph("• Peso proprio della muratura", style='List Bullet')
        doc.add_paragraph("• Carichi permanenti dai solai", style='List Bullet')
        doc.add_paragraph("• Carichi variabili di esercizio", style='List Bullet')
        
        # Formula carico totale
        self._add_formula_box(doc,
            "N_Ed = G_1 + G_2 + ψ_0 × Q_k\n" +
            "dove:\n" +
            "G_1 = carichi permanenti strutturali\n" +
            "G_2 = carichi permanenti non strutturali\n" +
            "Q_k = carichi variabili caratteristici\n" +
            "ψ_0 = coefficiente di combinazione"
        )
        
        p = doc.add_paragraph()
        p.add_run(f"Carico verticale totale: N_Ed = {loads.get('vertical', 150)} kN")
        
        doc.add_heading('6.2 Eccentricità dei carichi', 2)
        
        p = doc.add_paragraph()
        p.add_run("L'eccentricità del carico è data da:")
        
        self._add_formula_box(doc,
            "e = e_s + e_a\n" +
            "dove:\n" +
            "e_s = eccentricità strutturale\n" +
            "e_a = eccentricità accidentale = h/200"
        )
        
        p = doc.add_paragraph()
        p.add_run(f"Eccentricità totale: e = {loads.get('eccentricity', 10)} cm")
        
    def _add_calculation_procedure(self, doc):
        """Aggiunge procedura di calcolo dettagliata con formule"""
        doc.add_heading('7. PROCEDURA DI CALCOLO', 1)
        
        doc.add_heading('7.1 Metodo di verifica', 2)
        
        p = doc.add_paragraph()
        p.add_run("La verifica viene condotta secondo il metodo semiprobabilistico agli stati limite, ")
        p.add_run("confrontando lo stato tensionale con le resistenze di progetto dei materiali.")
        
        doc.add_heading('7.2 Rigidezza della parete', 2)
        
        p = doc.add_paragraph()
        p.add_run("La rigidezza laterale della parete muraria è calcolata secondo:")
        
        self._add_formula_box(doc,
            "K = (E × t) / [12 × (h/L)³ × (1 + 3φ)]\n" +
            "dove:\n" +
            "φ = 1.2 × (h/L)² × (E/G)\n" +
            "E = modulo elastico normale\n" +
            "G = modulo elastico tangenziale\n" +
            "t = spessore della parete\n" +
            "h = altezza della parete\n" +
            "L = lunghezza della parete"
        )
        
        doc.add_heading('7.3 Resistenza a taglio', 2)
        
        p = doc.add_paragraph()
        p.add_run("La resistenza a taglio è valutata secondo tre meccanismi:")
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Meccanismo 1 - Rottura per taglio puro:").bold = True
        
        self._add_formula_box(doc,
            "V_t1 = (L × t × f_vd) / FC\n" +
            "dove:\n" +
            "f_vd = f_vk / γ_M = resistenza di progetto a taglio\n" +
            "FC = fattore di confidenza"
        )
        
        p = doc.add_paragraph()
        p.add_run("Meccanismo 2 - Rottura per taglio con attrito:").bold = True
        
        self._add_formula_box(doc,
            "V_t2 = (L × t × τ_0d + μ × σ_n) / FC\n" +
            "dove:\n" +
            "τ_0d = resistenza a taglio in assenza di carichi verticali\n" +
            "μ = 0.4 = coefficiente di attrito\n" +
            "σ_n = N / (L × t) = tensione normale media"
        )
        
        p = doc.add_paragraph()
        p.add_run("Meccanismo 3 - Rottura per pressoflessione:").bold = True
        
        self._add_formula_box(doc,
            "V_t3 = (L × t × f_cd / 2) × (1 - σ_n / f_cd) × √(1 - (2e/t)²)\n" +
            "dove:\n" +
            "f_cd = resistenza di progetto a compressione\n" +
            "e = eccentricità del carico"
        )
        
    def _add_existing_state_verification(self, doc):
        """Aggiunge verifiche stato di fatto con calcoli"""
        doc.add_heading('8. VERIFICHE STATO DI FATTO', 1)
        
        orig = self.results.get('original', {})
        
        doc.add_heading('8.1 Calcolo della rigidezza', 2)
        
        wall = self.project_data.get('wall', {})
        masonry = self.project_data.get('masonry', {})
        
        # Parametri
        L = wall.get('length', 500) / 100  # in metri
        h = wall.get('height', 300) / 100
        t = wall.get('thickness', 50) / 100
        E = masonry.get('E', 1740) * 1000  # in kN/m²
        G = E / 6
        
        # Calcoli step by step
        p = doc.add_paragraph()
        p.add_run("Parametri geometrici e meccanici:")
        
        params = [
            f"L = {L} m",
            f"h = {h} m", 
            f"t = {t} m",
            f"E = {E/1000} MPa",
            f"G = {G/1000:.0f} MPa"
        ]
        
        for param in params:
            doc.add_paragraph(f"• {param}", style='List Bullet')
            
        p = doc.add_paragraph()
        p.add_run("Calcolo del parametro φ:")
        
        phi = 1.2 * (h/L)**2 * (E/G)
        self._add_formula_box(doc,
            f"φ = 1.2 × ({h}/{L})² × ({E/1000}/{G/1000:.0f}) = {phi:.3f}"
        )
        
        p = doc.add_paragraph()
        p.add_run("Calcolo della rigidezza:")
        
        K = orig.get('K', 0)
        self._add_formula_box(doc,
            f"K = ({E/1000} × {t}) / [12 × ({h}/{L})³ × (1 + 3×{phi:.3f})]\n" +
            f"K = {K:.1f} kN/m"
        )
        
        doc.add_heading('8.2 Calcolo delle resistenze', 2)
        
        # Resistenza V_t1
        p = doc.add_paragraph()
        p.add_run("Resistenza per taglio puro (V_t1):")
        
        V_t1 = orig.get('V_t1', 0)
        FC = self.results.get('FC', 1.35)
        tau0 = masonry.get('tau0', 0.076)
        
        self._add_formula_box(doc,
            f"V_t1 = (L × t × τ_0d) / FC\n" +
            f"V_t1 = ({L} × {t} × {tau0*1000}) / {FC}\n" +
            f"V_t1 = {V_t1:.1f} kN"
        )
        
        # Resistenza V_t2
        p = doc.add_paragraph()
        p.add_run("Resistenza per taglio con attrito (V_t2):")
        
        V_t2 = orig.get('V_t2', 0)
        N = self.project_data.get('loads', {}).get('vertical', 150)
        
        self._add_formula_box(doc,
            f"V_t2 = (L × t × τ_0d + 0.4 × N) / FC\n" +
            f"V_t2 = ({L} × {t} × {tau0*1000} + 0.4 × {N}) / {FC}\n" +
            f"V_t2 = {V_t2:.1f} kN"
        )
        
        # Riepilogo risultati
        doc.add_heading('8.3 Riepilogo risultati stato di fatto', 2)
        
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        results_data = [
            ('Parametro', 'Valore'),
            ('Rigidezza laterale K', f"{K:.1f} kN/m"),
            ('Resistenza V_t1', f"{V_t1:.1f} kN"),
            ('Resistenza V_t2', f"{V_t2:.1f} kN"),
            ('Resistenza minima V_min', f"{orig.get('V_min', 0):.1f} kN")
        ]
        
        for i, (param, value) in enumerate(results_data):
            table.rows[i].cells[0].text = param
            table.rows[i].cells[1].text = value
            
    def _add_project_state_verification(self, doc):
        """Aggiunge verifiche stato di progetto"""
        doc.add_heading('9. VERIFICHE STATO DI PROGETTO', 1)
        
        mod = self.results.get('modified', {})
        
        doc.add_heading('9.1 Effetto delle aperture', 2)
        
        p = doc.add_paragraph()
        p.add_run("Le aperture riducono la sezione resistente della parete. ")
        p.add_run("La nuova rigidezza è calcolata considerando:")
        
        items = [
            "Riduzione della lunghezza efficace della parete",
            "Contributo delle cerchiature metalliche",
            "Modifica dello schema statico"
        ]
        
        # Aggiungi nota per aperture ad arco
        openings = self.project_data.get('openings_module', {}).get('openings', [])
        n_arch = sum(1 for op in openings if op.get('type', 'Rettangolare') == 'Ad arco')
        if n_arch > 0:
            items.append("Geometria curvilinea delle aperture ad arco")
        
        for item in items:
            doc.add_paragraph(f"• {item}", style='List Bullet')
            
        doc.add_heading('9.2 Contributo delle cerchiature', 2)
        
        p = doc.add_paragraph()
        p.add_run("Il contributo alla rigidezza delle cerchiature metalliche è valutato come:")
        
        self._add_formula_box(doc,
            "K_cerchiature = Σ (12 × E_s × I_eq) / h³\n" +
            "dove:\n" +
            "E_s = modulo elastico acciaio\n" +
            "I_eq = momento d'inerzia equivalente del telaio\n" +
            "h = altezza dell'apertura"
        )
        
        if n_arch > 0:
            p = doc.add_paragraph()
            p.add_run("Per le aperture ad arco, il momento d'inerzia equivalente tiene conto della ")
            p.add_run("geometria curvilinea e della eventuale calandratura dei profili.")
        
        K_cerch = mod.get('K_cerchiature', 0)
        p = doc.add_paragraph()
        p.add_run(f"Contributo cerchiature: K_cerchiature = {K_cerch:.1f} kN/m")
        
        doc.add_heading('9.3 Riepilogo risultati stato di progetto', 2)
        
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        
        results_data = [
            ('Parametro', 'Valore'),
            ('Rigidezza laterale K', f"{mod.get('K', 0):.1f} kN/m"),
            ('Contributo cerchiature', f"{K_cerch:.1f} kN/m"),
            ('Resistenza V_t1', f"{mod.get('V_t1', 0):.1f} kN"),
            ('Resistenza V_t2', f"{mod.get('V_t2', 0):.1f} kN"),
            ('Resistenza minima V_min', f"{mod.get('V_min', 0):.1f} kN")
        ]
        
        for i, (param, value) in enumerate(results_data):
            table.rows[i].cells[0].text = param
            table.rows[i].cells[1].text = value
            
    def _add_normative_verifications(self, doc):
        """Aggiunge verifiche normative con diagrammi"""
        doc.add_heading('10. CONFRONTO E VERIFICHE NORMATIVE', 1)
        
        verif = self.results.get('verification', {})
        
        doc.add_heading('10.1 Verifica della variazione di rigidezza', 2)
        
        p = doc.add_paragraph()
        p.add_run("Secondo NTC 2018 § 8.4.1, la variazione di rigidezza deve essere:")
        
        self._add_formula_box(doc,
            "|ΔK| = |K_progetto - K_originale| / K_originale × 100 ≤ 15%"
        )
        
        stiff_var = verif.get('stiffness_variation', 55.9)
        stiff_ok = verif.get('stiffness_ok', False)
        
        p = doc.add_paragraph()
        p.add_run(f"Variazione calcolata: ΔK = {stiff_var:.1f}%")
        
        p = doc.add_paragraph()
        p.add_run("Esito verifica: ")
        if stiff_ok:
            p.add_run("VERIFICATA ✓").font.color.rgb = RGBColor(0, 128, 0)
        else:
            p.add_run("NON VERIFICATA ✗").font.color.rgb = RGBColor(255, 0, 0)
            
        doc.add_heading('10.2 Verifica della variazione di resistenza', 2)
        
        p = doc.add_paragraph()
        p.add_run("La riduzione di resistenza deve essere:")
        
        self._add_formula_box(doc,
            "ΔV = (V_originale - V_progetto) / V_originale × 100 ≤ 20%"
        )
        
        res_var = verif.get('resistance_variation', 28.4)
        res_ok = verif.get('resistance_ok', False)
        
        p = doc.add_paragraph()
        p.add_run(f"Variazione calcolata: ΔV = {res_var:.1f}%")
        
        p = doc.add_paragraph()
        p.add_run("Esito verifica: ")
        if res_ok:
            p.add_run("VERIFICATA ✓").font.color.rgb = RGBColor(0, 128, 0)
        else:
            p.add_run("NON VERIFICATA ✗").font.color.rgb = RGBColor(255, 0, 0)
            
        # Inserisci diagramma comparativo
        if MATPLOTLIB_AVAILABLE:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Figura 10.1 - Confronto rigidezze e resistenze").italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            img_path = self._generate_comparison_chart()
            if img_path and os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(5))
                
    def _add_conclusions(self, doc):
        """Aggiunge conclusioni"""
        doc.add_heading('11. CONCLUSIONI', 1)
        
        verif = self.results.get('verification', {})
        is_local = verif.get('is_local', False)
        
        p = doc.add_paragraph()
        p.add_run("Sulla base delle verifiche effettuate secondo le NTC 2018, si conclude quanto segue:")
        
        doc.add_paragraph()
        
        if is_local:
            p = doc.add_paragraph()
            p.add_run("L'intervento è classificabile come ").bold = False
            p.add_run("INTERVENTO LOCALE").bold = True
            p.add_run(" ai sensi del § 8.4.1 delle NTC 2018.")
            
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Pertanto non è necessaria la valutazione della sicurezza globale della costruzione.")
        else:
            p = doc.add_paragraph()
            p.add_run("L'intervento ").bold = False
            p.add_run("NON").bold = True
            p.add_run(" è classificabile come intervento locale.")
            
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Le variazioni di rigidezza e/o resistenza superano i limiti normativi:")
            
            issues = []
            if verif.get('stiffness_variation', 55.9) > 15:
                issues.append(f"Variazione di rigidezza: {verif.get('stiffness_variation', 55.9):.1f}% > 15%")
            if verif.get('resistance_variation', 28.4) > 20:
                issues.append(f"Variazione di resistenza: {verif.get('resistance_variation', 28.4):.1f}% > 20%")
                
            for issue in issues:
                doc.add_paragraph(f"• {issue}", style='List Bullet')
                
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Si rendono necessarie le seguenti azioni:").bold = True
            
            actions = [
                "Ridimensionamento delle aperture per ridurre l'impatto strutturale",
                "Inserimento di rinforzi aggiuntivi (FRP, intonaco armato, etc.)",
                "Valutazione della sicurezza globale dell'edificio",
                "Eventuale classificazione come intervento di miglioramento sismico"
            ]
            
            for action in actions:
                doc.add_paragraph(f"• {action}", style='List Bullet')
                
        # Firma
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Luogo e data
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"Palermo, {datetime.now().strftime('%d/%m/%Y')}")
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Timbro e firma
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run("Il Progettista Strutturale")
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(self.engineer_data['nome']).bold = True
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(self.engineer_data['iscrizione']).italic = True
        
    def _add_attachments(self, doc):
        """Aggiunge sezione allegati"""
        doc.add_heading('ALLEGATI', 1)
        
        doc.add_heading('A. DISEGNI TECNICI', 2)
        
        attachments = [
            "A.1 - Pianta stato di fatto",
            "A.2 - Pianta stato di progetto", 
            "A.3 - Sezione verticale con dettaglio cerchiature",
            "A.4 - Particolari costruttivi cerchiature"
        ]
        
        # Aggiungi allegati specifici per aperture ad arco
        openings = self.project_data.get('openings_module', {}).get('openings', [])
        n_arch = sum(1 for op in openings if op.get('type', 'Rettangolare') == 'Ad arco')
        if n_arch > 0:
            attachments.append("A.5 - Dettaglio calandratura profili per aperture ad arco")
            attachments.append("A.6 - Schema delle fasi esecutive per aperture curvilinee")
        else:
            attachments.append("A.5 - Schema delle fasi esecutive")
        
        for att in attachments:
            doc.add_paragraph(att, style='List Bullet')
            
        doc.add_heading('B. DIAGRAMMI DI CALCOLO', 2)
        
        diagrams = [
            "B.1 - Diagramma delle rigidezze",
            "B.2 - Diagramma delle resistenze",
            "B.3 - Confronto stato di fatto/progetto",
            "B.4 - Distribuzione delle tensioni"
        ]
        
        for diag in diagrams:
            doc.add_paragraph(diag, style='List Bullet')
            
        doc.add_heading('C. TABULATI DI CALCOLO', 2)
        
        p = doc.add_paragraph()
        p.add_run("Di seguito i tabulati completi dei calcoli effettuati:")
        
        # Tabella riassuntiva calcoli
        self._add_calculation_summary_table(doc)
        
    def _add_formula_box(self, doc, formula_text):
        """Aggiunge un box per le formule"""
        # Crea tabella con bordo per formula
        table = doc.add_table(rows=1, cols=1)
        table.style = None
        
        cell = table.rows[0].cells[0]
        cell._element.get_or_add_tcPr().append(self._create_shading_element())
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Aggiungi testo formula con font monospace
        for line in formula_text.split('\n'):
            if line:
                run = p.add_run(line)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                if p.runs:
                    p.add_run('\n')
                    
    def _create_shading_element(self):
        """Crea elemento di sfondo per celle"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'F0F0F0')
        return shading_elm
        
    def _add_calculation_summary_table(self, doc):
        """Aggiunge tabella riassuntiva dei calcoli"""
        table = doc.add_table(rows=12, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Intestazioni
        headers = ['Parametro', 'Stato di Fatto', 'Stato di Progetto', 'Variazione']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            
        # Dati
        orig = self.results.get('original', {})
        mod = self.results.get('modified', {})
        verif = self.results.get('verification', {})
        
        data = [
            ('GEOMETRIA', '', '', ''),
            ('Area resistente [m²]', f"{orig.get('area', 0):.2f}", f"{mod.get('area', 0):.2f}", 
             f"{((mod.get('area', 0) - orig.get('area', 0))/orig.get('area', 1)*100):.1f}%"),
            ('RIGIDEZZA', '', '', ''),
            ('K muratura [kN/m]', f"{orig.get('K', 0):.1f}", f"{mod.get('K', 0) - mod.get('K_cerchiature', 0):.1f}", '-'),
            ('K cerchiature [kN/m]', '0', f"{mod.get('K_cerchiature', 0):.1f}", '-'),
            ('K totale [kN/m]', f"{orig.get('K', 0):.1f}", f"{mod.get('K', 0):.1f}", 
             f"{verif.get('stiffness_variation', 0):.1f}%"),
            ('RESISTENZA', '', '', ''),
            ('V_t1 [kN]', f"{orig.get('V_t1', 0):.1f}", f"{mod.get('V_t1', 0):.1f}", '-'),
            ('V_t2 [kN]', f"{orig.get('V_t2', 0):.1f}", f"{mod.get('V_t2', 0):.1f}", '-'),
            ('V_min [kN]', f"{orig.get('V_min', 0):.1f}", f"{mod.get('V_min', 0):.1f}", 
             f"{verif.get('resistance_variation', 0):.1f}%"),
            ('VERIFICHE', '', '', '')
        ]
        
        for i, (param, sf, sp, var) in enumerate(data, start=1):
            table.rows[i].cells[0].text = param
            table.rows[i].cells[1].text = sf
            table.rows[i].cells[2].text = sp
            table.rows[i].cells[3].text = var
            
            # Evidenzia righe di sezione
            if param in ['GEOMETRIA', 'RIGIDEZZA', 'RESISTENZA', 'VERIFICHE']:
                for cell in table.rows[i].cells:
                    cell._element.get_or_add_tcPr().append(self._create_shading_element())
                    
    def _generate_wall_drawing(self, state='existing'):
        """Genera disegno tecnico della parete con supporto aperture ad arco"""
        if not MATPLOTLIB_AVAILABLE:
            return None
            
        try:
            from matplotlib.path import Path
            from matplotlib.patches import PathPatch, Wedge
            import numpy as np
            
            fig, ax = plt.subplots(figsize=(10, 6))
            
            wall = self.project_data.get('wall', {})
            L = wall.get('length', 500)
            h = wall.get('height', 300)
            t = wall.get('thickness', 50)
            
            # Disegna parete
            wall_rect = Rectangle((0, 0), L, h, 
                                linewidth=2, 
                                edgecolor='black', 
                                facecolor='lightgray')
            ax.add_patch(wall_rect)
            
            # Se stato di progetto, aggiungi aperture
            if state == 'project':
                openings = self.project_data.get('openings_module', {}).get('openings', [])
                for i, opening in enumerate(openings):
                    x = opening.get('x', 0)
                    w = opening.get('width', 120)
                    h_op = opening.get('height', 240)
                    base = opening.get('base_height', 0)
                    op_type = opening.get('type', 'Rettangolare')
                    
                    if op_type == 'Ad arco':
                        # Parametri dell'arco
                        arch_params = opening.get('arch_params', {})
                        rise = arch_params.get('rise', h_op * 0.2)  # Freccia dell'arco
                        rect_height = h_op - rise
                        
                        # Crea un percorso completo per l'apertura ad arco
                        # che include sia la parte rettangolare che l'arco
                        
                        # Numero di punti per l'arco
                        n_arc_points = 30
                        arc_angles = np.linspace(0, np.pi, n_arc_points)
                        arc_center_x = x + w/2
                        arc_center_y = base + rect_height
                        arc_radius_x = w/2
                        arc_radius_y = rise
                        
                        # Costruisci il percorso dell'apertura
                        verts = []
                        codes = []
                        
                        # Parte inferiore sinistra
                        verts.append((x, base))
                        codes.append(Path.MOVETO)
                        
                        # Lato sinistro fino all'imposta dell'arco
                        verts.append((x, base + rect_height))
                        codes.append(Path.LINETO)
                        
                        # Arco superiore (da sinistra a destra)
                        for angle in arc_angles:
                            arc_x = arc_center_x - arc_radius_x * np.cos(angle)
                            arc_y = arc_center_y + arc_radius_y * np.sin(angle)
                            verts.append((arc_x, arc_y))
                            codes.append(Path.LINETO)
                        
                        # Lato destro dall'imposta dell'arco in giù
                        verts.append((x + w, base))
                        codes.append(Path.LINETO)
                        
                        # Chiudi il percorso
                        verts.append((x, base))
                        codes.append(Path.CLOSEPOLY)
                        
                        # Crea il percorso e riempilo di bianco
                        path = Path(verts, codes)
                        opening_patch = PathPatch(path, 
                                                facecolor='white', 
                                                edgecolor='red', 
                                                linewidth=2,
                                                zorder=2)  # Assicura che sia sopra la muratura
                        ax.add_patch(opening_patch)
                        
                        # Disegna le cerchiature
                        # Cerchiatura curva esterna (solo contorno)
                        frame_arc = Arc((arc_center_x, arc_center_y), 
                                       w + 20, 
                                       2 * rise + 20,
                                       angle=0, theta1=0, theta2=180,
                                       linewidth=3, 
                                       edgecolor='blue', 
                                       linestyle='--',
                                       zorder=3)
                        ax.add_artist(frame_arc)
                        
                        # Piedritti laterali della cerchiatura
                        left_frame = Rectangle((x-10, base-10), 10, rect_height+10,
                                             linewidth=3,
                                             edgecolor='blue',
                                             facecolor='none',
                                             linestyle='--',
                                             zorder=3)
                        ax.add_patch(left_frame)
                        
                        right_frame = Rectangle((x+w, base-10), 10, rect_height+10,
                                              linewidth=3,
                                              edgecolor='blue',
                                              facecolor='none',
                                              linestyle='--',
                                              zorder=3)
                        ax.add_patch(right_frame)
                        
                        # Base della cerchiatura
                        bottom_frame = Rectangle((x-10, base-10), w+20, 10,
                                               linewidth=3,
                                               edgecolor='blue',
                                               facecolor='none',
                                               linestyle='--',
                                               zorder=3)
                        ax.add_patch(bottom_frame)
                        
                        # Etichetta
                        ax.text(x + w/2, base + h_op/2, f'AP{i+1}\n(ARCO)',
                               ha='center', va='center', 
                               fontsize=12, fontweight='bold',
                               zorder=4)
                               
                    else:
                        # Apertura rettangolare standard
                        opening_rect = Rectangle((x, base), w, h_op,
                                              linewidth=2,
                                              edgecolor='red',
                                              facecolor='white',
                                              zorder=2)
                        ax.add_patch(opening_rect)
                        
                        # Cerchiatura rettangolare
                        frame_rect = Rectangle((x-10, base-10), w+20, h_op+20,
                                             linewidth=3,
                                             edgecolor='blue',
                                             facecolor='none',
                                             linestyle='--',
                                             zorder=3)
                        ax.add_patch(frame_rect)
                        
                        # Etichetta
                        ax.text(x + w/2, base + h_op/2, f'AP{i+1}',
                               ha='center', va='center', 
                               fontsize=12, fontweight='bold',
                               zorder=4)
                    
            # Quote
            ax.plot([0, L], [-50, -50], 'k-', linewidth=1)
            ax.plot([0, 0], [-40, -60], 'k-', linewidth=1)
            ax.plot([L, L], [-40, -60], 'k-', linewidth=1)
            ax.text(L/2, -80, f'{L} cm', ha='center', fontsize=10)
            
            ax.plot([-50, -50], [0, h], 'k-', linewidth=1)
            ax.plot([-40, -60], [0, 0], 'k-', linewidth=1)
            ax.plot([-40, -60], [h, h], 'k-', linewidth=1)
            ax.text(-100, h/2, f'{h} cm', ha='center', va='center', rotation=90, fontsize=10)
            
            # Impostazioni
            ax.set_xlim(-150, L+50)
            ax.set_ylim(-100, h+50)
            ax.set_aspect('equal')
            ax.axis('off')
            
            # Titolo
            title = "STATO DI FATTO" if state == 'existing' else "STATO DI PROGETTO"
            ax.text(L/2, h+30, title, ha='center', fontsize=14, fontweight='bold')
            
            # Legenda se stato di progetto
            if state == 'project':
                # Aggiungi legenda
                legend_elements = [
                    mpatches.Rectangle((0, 0), 1, 1, facecolor='lightgray', 
                                     edgecolor='black', label='Muratura esistente'),
                    mpatches.Rectangle((0, 0), 1, 1, facecolor='white', 
                                     edgecolor='red', label='Apertura'),
                    mpatches.Rectangle((0, 0), 1, 1, facecolor='none', 
                                     edgecolor='blue', linestyle='--', label='Cerchiatura')
                ]
                ax.legend(handles=legend_elements, loc='upper right', bbox_to_anchor=(1.1, 1))
            
            # Salva
            filename = f'wall_drawing_{state}.png'
            plt.savefig(filename, dpi=150, bbox_inches='tight')
            plt.close()
            
            return filename
            
        except Exception as e:
            print(f"Errore generazione disegno: {e}")
            return None
            
    def _generate_comparison_chart(self):
        """Genera grafico comparativo"""
        if not MATPLOTLIB_AVAILABLE:
            return None
            
        try:
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 6))
            
            orig = self.results.get('original', {})
            mod = self.results.get('modified', {})
            
            # Grafico rigidezze
            categories = ['Stato di\nFatto', 'Stato di\nProgetto']
            rigidezze = [orig.get('K', 0), mod.get('K', 0)]
            
            bars1 = ax1.bar(categories, rigidezze, color=['green', 'orange'])
            ax1.set_ylabel('Rigidezza [kN/m]')
            ax1.set_title('Confronto Rigidezze')
            
            # Aggiungi valori sulle barre
            for bar, val in zip(bars1, rigidezze):
                height = bar.get_height()
                ax1.text(bar.get_x() + bar.get_width()/2., height,
                        f'{val:.1f}', ha='center', va='bottom')
                        
            # Linea limite variazione
            ax1.axhline(y=orig.get('K', 0)*0.85, color='red', linestyle='--', 
                       label='Limite -15%')
            ax1.axhline(y=orig.get('K', 0)*1.15, color='red', linestyle='--',
                       label='Limite +15%')
            ax1.legend()
            
            # Grafico resistenze
            resistenze = [orig.get('V_min', 0), mod.get('V_min', 0)]
            
            bars2 = ax2.bar(categories, resistenze, color=['green', 'orange'])
            ax2.set_ylabel('Resistenza [kN]')
            ax2.set_title('Confronto Resistenze')
            
            # Aggiungi valori
            for bar, val in zip(bars2, resistenze):
                height = bar.get_height()
                ax2.text(bar.get_x() + bar.get_width()/2., height,
                        f'{val:.1f}', ha='center', va='bottom')
                        
            # Linea limite riduzione
            ax2.axhline(y=orig.get('V_min', 0)*0.8, color='red', linestyle='--',
                       label='Limite -20%')
            ax2.legend()
            
            plt.suptitle('CONFRONTO STATO DI FATTO - STATO DI PROGETTO', 
                        fontsize=14, fontweight='bold')
            
            plt.tight_layout()
            
            filename = 'comparison_chart.png'
            plt.savefig(filename, dpi=150)
            plt.close()
            
            return filename
            
        except Exception as e:
            print(f"Errore generazione grafico: {e}")
            return None
            

# Aggiorna il modulo principale per usare il generatore avanzato
class EnhancedReportModule(QWidget):
    """Modulo interfaccia per generazione relazioni avanzate"""
    
    data_changed = pyqtSignal()
    
    def __init__(self):
        super().__init__()
        self.project_data = {}
        self.results = {}
        self.generator = EnhancedReportGenerator()
        self.setup_ui()
        
    def setup_ui(self):
        """Costruisce interfaccia"""
        layout = QVBoxLayout()
        
        # Gruppo opzioni
        options_group = QGroupBox("Opzioni Relazione Avanzata")
        options_layout = QVBoxLayout()
        
        # Tipo relazione
        report_type_layout = QHBoxLayout()
        report_type_layout.addWidget(QLabel("Tipo relazione:"))
        
        self.standard_radio = QRadioButton("Standard")
        self.complete_radio = QRadioButton("Completa con formule")
        self.complete_radio.setChecked(True)
        
        report_type_layout.addWidget(self.standard_radio)
        report_type_layout.addWidget(self.complete_radio)
        report_type_layout.addStretch()
        options_layout.addLayout(report_type_layout)
        
        # Contenuti da includere
        self.include_formulas = QCheckBox("Includi formule dettagliate")
        self.include_formulas.setChecked(True)
        options_layout.addWidget(self.include_formulas)
        
        self.include_drawings = QCheckBox("Includi disegni tecnici stato fatto/progetto")
        self.include_drawings.setChecked(True)
        options_layout.addWidget(self.include_drawings)
        
        self.include_charts = QCheckBox("Includi diagrammi comparativi")
        self.include_charts.setChecked(True)
        options_layout.addWidget(self.include_charts)
        
        self.include_photos = QCheckBox("Spazio per documentazione fotografica")
        options_layout.addWidget(self.include_photos)
        
        self.include_attachments = QCheckBox("Includi sezione allegati")
        self.include_attachments.setChecked(True)
        options_layout.addWidget(self.include_attachments)
        
        options_group.setLayout(options_layout)
        layout.addWidget(options_group)
        
        # Anteprima
        preview_group = QGroupBox("Anteprima Struttura")
        preview_layout = QVBoxLayout()
        
        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)
        preview_layout.addWidget(self.preview_text)
        
        preview_group.setLayout(preview_layout)
        layout.addWidget(preview_group)
        
        # Pulsanti
        button_layout = QHBoxLayout()
        
        self.preview_btn = QPushButton("Aggiorna Anteprima")
        self.preview_btn.clicked.connect(self.update_preview)
        button_layout.addWidget(self.preview_btn)
        
        self.generate_btn = QPushButton("Genera Relazione Completa")
        self.generate_btn.clicked.connect(self.generate_report)
        self.generate_btn.setEnabled(DOCX_AVAILABLE)
        button_layout.addWidget(self.generate_btn)
        
        layout.addLayout(button_layout)
        
        # Avvisi
        if not DOCX_AVAILABLE:
            warning = QLabel("⚠️ Installare python-docx: pip install python-docx")
            warning.setStyleSheet("color: red;")
            layout.addWidget(warning)
            
        if not MATPLOTLIB_AVAILABLE:
            warning = QLabel("⚠️ Per i grafici installare: pip install matplotlib")
            warning.setStyleSheet("color: orange;")
            layout.addWidget(warning)
        
        self.setLayout(layout)
        
    def update_preview(self):
        """Aggiorna anteprima struttura relazione"""
        preview = []
        
        preview.append("STRUTTURA RELAZIONE TECNICA COMPLETA")
        preview.append("=" * 40)
        preview.append("")
        
        sections = [
            "1. PREMESSA",
            "2. RIFERIMENTI NORMATIVI",
            "   2.1 Normativa nazionale",
            "   2.2 Riferimenti specifici NTC 2018",
            "3. DESCRIZIONE DELL'INTERVENTO",
            "   3.1 Descrizione generale",
            "   3.2 Classificazione dell'intervento",
            "4. CARATTERISTICHE DEI MATERIALI",
            "   4.1 Muratura esistente",
            "   4.2 Acciaio per cerchiature",
            "5. GEOMETRIA DELLA STRUTTURA",
            "   5.1 Geometria della parete",
            "   5.2 Aperture previste"
        ]
        
        if self.include_drawings.isChecked():
            sections.extend([
                "   - Disegno stato di fatto",
                "   - Disegno stato di progetto"
            ])
            
        sections.extend([
            "6. ANALISI DEI CARICHI",
            "   6.1 Carichi verticali",
            "   6.2 Eccentricità dei carichi",
            "7. PROCEDURA DI CALCOLO"
        ])
        
        if self.include_formulas.isChecked():
            sections.extend([
                "   7.1 Metodo di verifica",
                "   7.2 Rigidezza della parete (con formule)",
                "   7.3 Resistenza a taglio (con formule)"
            ])
            
        sections.extend([
            "8. VERIFICHE STATO DI FATTO",
            "   8.1 Calcolo della rigidezza",
            "   8.2 Calcolo delle resistenze",
            "   8.3 Riepilogo risultati",
            "9. VERIFICHE STATO DI PROGETTO",
            "   9.1 Effetto delle aperture",
            "   9.2 Contributo delle cerchiature",
            "   9.3 Riepilogo risultati",
            "10. CONFRONTO E VERIFICHE NORMATIVE",
            "   10.1 Verifica variazione rigidezza",
            "   10.2 Verifica variazione resistenza"
        ])
        
        if self.include_charts.isChecked():
            sections.append("   - Diagrammi comparativi")
            
        sections.extend([
            "11. CONCLUSIONI",
        ])
        
        if self.include_attachments.isChecked():
            sections.extend([
                "",
                "ALLEGATI",
                "A. DISEGNI TECNICI",
                "B. DIAGRAMMI DI CALCOLO",
                "C. TABULATI DI CALCOLO"
            ])
            
        if self.include_photos.isChecked():
            sections.append("D. DOCUMENTAZIONE FOTOGRAFICA")
            
        # Risultati
        preview.extend(sections)
        preview.append("")
        preview.append("=" * 40)
        
        if self.results:
            verif = self.results.get('verification', {})
            preview.append("")
            preview.append("ESITO VERIFICHE:")
            if verif.get('is_local', False):
                preview.append("✓ INTERVENTO LOCALE")
            else:
                preview.append("✗ INTERVENTO NON LOCALE")
                preview.append(f"  - Var. rigidezza: {verif.get('stiffness_variation', 0):.1f}%")
                preview.append(f"  - Var. resistenza: {verif.get('resistance_variation', 0):.1f}%")
                
        self.preview_text.setPlainText('\n'.join(preview))
        
    def set_data(self, project_data: Dict, results: Dict):
        """Imposta dati progetto e risultati"""
        self.project_data = project_data
        self.results = results
        self.generator.set_data(project_data, results)
        self.update_preview()
        
    def generate_report(self):
        """Genera relazione completa"""
        if not self.project_data or not self.results:
            QMessageBox.warning(self, "Attenzione", 
                              "Dati mancanti. Eseguire prima il calcolo.")
            return
            
        # Dialog salvataggio
        filename, _ = QFileDialog.getSaveFileName(
            self, "Salva Relazione Tecnica Completa",
            f"Relazione_Tecnica_Cerchiature_{datetime.now().strftime('%Y%m%d')}.docx",
            "Documenti Word (*.docx)"
        )
        
        if filename:
            try:
                # Configura generatore
                self.generator.include_formulas = self.include_formulas.isChecked()
                self.generator.include_drawings = self.include_drawings.isChecked()
                self.generator.include_charts = self.include_charts.isChecked()
                self.generator.include_attachments = self.include_attachments.isChecked()
                
                # Genera relazione
                self.generator.generate_enhanced_word_report(filename)
                
                QMessageBox.information(self, "Completato",
                                      f"Relazione tecnica completa generata:\n{filename}")
                                      
                # Apri file
                if os.path.exists(filename):
                    os.startfile(filename)
                    
            except Exception as e:
                QMessageBox.critical(self, "Errore",
                                   f"Errore generazione relazione:\n{str(e)}")