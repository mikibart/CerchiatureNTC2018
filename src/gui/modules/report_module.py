"""
Modulo Generazione Relazione Tecnica
Genera relazioni tecniche professionali per interventi di cerchiatura
Arch. Michelangelo Bartolotta
"""

from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtPrintSupport import QPrinter, QPrintDialog
from datetime import datetime
import os
from typing import Dict
import json

# Per generazione PDF
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("reportlab non installato - generazione PDF non disponibile")

# Per generazione Word
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx non installato - generazione Word non disponibile")


class ReportGenerator:
    """Generatore di relazioni tecniche"""
    
    def __init__(self):
        self.project_data = {}
        self.results = {}
        self.engineer_data = {
            'nome': 'Arch. Michelangelo Bartolotta',
            'titolo': 'Architetto',
            'iscrizione': 'Ordine degli Architetti di Agrigento n. 1557',
            'cf': 'BRTMHL68C19A089V',
            'piva': '02554080842',
            'sede_op': 'Via Domenico Scinà n. 28, Palermo',
            'sede_legale': 'C.le S.to Argento n.11, 92020 Castrofilippo (AG)',
            'email': 'arch.bartolotta@pec.it',
            'tel': '+39 XXX XXXXXXX'
        }
        
    def set_data(self, project_data: Dict, results: Dict):
        """Imposta i dati del progetto e i risultati"""
        self.project_data = project_data
        self.results = results
        
    def generate_word_report(self, filename: str):
        """Genera relazione in formato Word"""
        if not DOCX_AVAILABLE:
            raise Exception("Modulo python-docx non disponibile")
            
        doc = Document()
        
        # Imposta margini
        section = doc.sections[0]
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        
        # Intestazione
        self._add_word_header(doc)
        
        # Titolo
        title = doc.add_heading('RELAZIONE TECNICA', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('VERIFICA LOCALE APERTURA IN PARETE MURARIA', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Riferimenti normativi
        doc.add_heading('1. RIFERIMENTI NORMATIVI', 1)
        self._add_word_normative(doc)
        
        # Descrizione intervento
        doc.add_heading('2. DESCRIZIONE DELL\'INTERVENTO', 1)
        self._add_word_description(doc)
        
        # Dati di input
        doc.add_heading('3. CARATTERISTICHE DEI MATERIALI', 1)
        self._add_word_materials(doc)
        
        # Geometria
        doc.add_heading('4. GEOMETRIA DELLA PARETE', 1)
        self._add_word_geometry(doc)
        
        # Analisi dei carichi
        doc.add_heading('5. ANALISI DEI CARICHI', 1)
        self._add_word_loads(doc)
        
        # Risultati calcolo
        doc.add_heading('6. RISULTATI DEL CALCOLO', 1)
        self._add_word_results(doc)
        
        # Verifiche
        doc.add_heading('7. VERIFICHE NORMATIVE', 1)
        self._add_word_verifications(doc)
        
        # Conclusioni
        doc.add_heading('8. CONCLUSIONI', 1)
        self._add_word_conclusions(doc)
        
        # Firma
        self._add_word_signature(doc)
        
        # Salva documento
        doc.save(filename)
        
    def generate_pdf_report(self, filename: str):
        """Genera relazione in formato PDF"""
        if not REPORTLAB_AVAILABLE:
            raise Exception("Modulo reportlab non disponibile")
            
        doc = SimpleDocTemplate(filename, pagesize=A4,
                              rightMargin=2.5*cm, leftMargin=2.5*cm,
                              topMargin=3*cm, bottomMargin=2.5*cm)
        
        # Stili
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='Center', alignment=TA_CENTER))
        styles.add(ParagraphStyle(name='Justify', alignment=TA_JUSTIFY))
        styles.add(ParagraphStyle(name='Right', alignment=TA_RIGHT))
        
        # Contenuto
        story = []
        
        # Intestazione
        self._add_pdf_header(story, styles)
        
        # Titolo
        story.append(Paragraph("RELAZIONE TECNICA", styles['Title']))
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph("VERIFICA LOCALE APERTURA IN PARETE MURARIA", styles['Heading1']))
        story.append(Spacer(1, 1*cm))
        
        # Sezioni
        self._add_pdf_normative(story, styles)
        self._add_pdf_description(story, styles)
        self._add_pdf_materials(story, styles)
        self._add_pdf_geometry(story, styles)
        self._add_pdf_loads(story, styles)
        self._add_pdf_results(story, styles)
        self._add_pdf_verifications(story, styles)
        self._add_pdf_conclusions(story, styles)
        self._add_pdf_signature(story, styles)
        
        # Genera PDF
        doc.build(story)
        
    def _add_word_header(self, doc):
        """Aggiunge intestazione professionale Word"""
        # Tabella per intestazione
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Cella sinistra - Dati professionista
        cell_left = table.rows[0].cells[0]
        p = cell_left.paragraphs[0]
        p.add_run(self.engineer_data['nome']).bold = True
        p.add_run(f"\n{self.engineer_data['titolo']}")
        p.add_run(f"\n{self.engineer_data['iscrizione']}")
        p.add_run(f"\n{self.engineer_data['sede_op']}")
        p.add_run(f"\nC.F. {self.engineer_data['cf']}")
        p.add_run(f"\nP.IVA {self.engineer_data['piva']}")
        
        # Cella destra - Dati progetto
        cell_right = table.rows[0].cells[1]
        p = cell_right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        info = self.project_data.get('info', {})
        p.add_run(f"Progetto: {info.get('name', 'N.D.')}")
        p.add_run(f"\nCommittente: {info.get('client', 'N.D.')}")
        p.add_run(f"\nUbicazione: {info.get('location', 'N.D.')}")
        p.add_run(f"\nData: {info.get('date', datetime.now().strftime('%d/%m/%Y'))}")
        
        doc.add_paragraph()
        
    def _add_word_normative(self, doc):
        """Aggiunge riferimenti normativi Word"""
        normative = [
            "• D.M. 17/01/2018 - Norme Tecniche per le Costruzioni (NTC 2018)",
            "• Circolare 21/01/2019 n. 7 - Istruzioni per l'applicazione delle NTC 2018",
            "• § 8.4.1 - Interventi locali o di riparazione",
            "• § 8.7.1 - Costruzioni in muratura",
            "• § 8.7.1.5 - Verifiche alle azioni non sismiche"
        ]
        
        for norm in normative:
            doc.add_paragraph(norm, style='List Bullet')
            
    def _add_word_description(self, doc):
        """Aggiunge descrizione intervento Word"""
        p = doc.add_paragraph()
        p.add_run("L'intervento in oggetto riguarda la realizzazione di ")
        
        # Conta aperture
        openings = self.project_data.get('openings_module', {}).get('openings', [])
        n_openings = len(openings)
        
        if n_openings == 0:
            p.add_run("aperture")
        elif n_openings == 1:
            p.add_run("n. 1 apertura")
        else:
            p.add_run(f"n. {n_openings} aperture")
            
        p.add_run(" in parete muraria portante, con inserimento di cerchiature metalliche.")
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Secondo NTC 2018 § 8.4.1, l'intervento è classificabile come ")
        p.add_run("LOCALE").bold = True
        p.add_run(" in quanto:")
        
        criteria = [
            "Non altera significativamente la rigidezza globale della struttura",
            "Non modifica il comportamento sismico complessivo",
            "Le variazioni di rigidezza e resistenza rimangono entro i limiti normativi"
        ]
        
        for criterion in criteria:
            doc.add_paragraph(f"• {criterion}", style='List Bullet')
            
    def _add_word_materials(self, doc):
        """Aggiunge caratteristiche materiali Word"""
        masonry = self.project_data.get('masonry', {})
        
        # Tabella materiali
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Intestazioni
        table.cell(0, 0).text = "Parametro"
        table.cell(0, 1).text = "Valore"
        
        # Dati
        data = [
            ("Tipologia muratura", masonry.get('type', 'N.D.')),
            ("Resistenza media compressione fcm", f"{masonry.get('fcm', 0)} MPa"),
            ("Resistenza media taglio τ0", f"{masonry.get('tau0', 0)} MPa"),
            ("Modulo elastico E", f"{masonry.get('E', 0)} MPa"),
            ("Livello di conoscenza", masonry.get('knowledge_level', 'LC1'))
        ]
        
        for i, (param, value) in enumerate(data, start=1):
            table.cell(i, 0).text = param
            table.cell(i, 1).text = value
            
        doc.add_paragraph()
        
        # Fattore di confidenza
        FC = self.results.get('FC', 1.35)
        p = doc.add_paragraph()
        p.add_run(f"Fattore di confidenza applicato: FC = {FC}")
        
    def _add_word_geometry(self, doc):
        """Aggiunge geometria parete Word"""
        wall = self.project_data.get('wall', {})
        
        p = doc.add_paragraph()
        p.add_run("Dimensioni della parete muraria:")
        
        dimensions = [
            f"Lunghezza: L = {wall.get('length', 0)} cm",
            f"Altezza: h = {wall.get('height', 0)} cm",
            f"Spessore: t = {wall.get('thickness', 0)} cm"
        ]
        
        for dim in dimensions:
            doc.add_paragraph(f"• {dim}", style='List Bullet')
            
        # Aperture
        openings = self.project_data.get('openings_module', {}).get('openings', [])
        if openings:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Aperture previste:")
            
            for i, opening in enumerate(openings, start=1):
                doc.add_paragraph(f"• Apertura {i}: {opening.get('width', 0)}×{opening.get('height', 0)} cm", 
                                style='List Bullet')
                
    def _add_word_loads(self, doc):
        """Aggiunge analisi carichi Word"""
        loads = self.project_data.get('loads', {})
        
        p = doc.add_paragraph()
        p.add_run("Carichi agenti sulla parete:")
        
        load_list = [
            f"Carico verticale: N = {loads.get('vertical', 0)} kN",
            f"Eccentricità: e = {loads.get('eccentricity', 0)} cm"
        ]
        
        for load in load_list:
            doc.add_paragraph(f"• {load}", style='List Bullet')
            
        # Vincoli
        constraints = self.project_data.get('constraints', {})
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Condizioni di vincolo:")
        doc.add_paragraph(f"• Base: {constraints.get('bottom', 'N.D.')}", style='List Bullet')
        doc.add_paragraph(f"• Sommità: {constraints.get('top', 'N.D.')}", style='List Bullet')
        
    def _add_word_results(self, doc):
        """Aggiunge risultati calcolo Word"""
        if not self.results:
            doc.add_paragraph("Risultati non disponibili")
            return
            
        # Stato di fatto
        doc.add_heading('6.1 Stato di Fatto', 2)
        
        orig = self.results.get('original', {})
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.cell(0, 0).text = "Parametro"
        table.cell(0, 1).text = "Valore"
        
        data = [
            ("Rigidezza laterale K", f"{orig.get('K', 0):.1f} kN/m"),
            ("Resistenza taglio V_t1", f"{orig.get('V_t1', 0):.1f} kN"),
            ("Resistenza taglio con f.f. V_t2", f"{orig.get('V_t2', 0):.1f} kN"),
            ("Resistenza pressoflessione V_t3", f"{orig.get('V_t3', 0):.1f} kN"),
            ("Resistenza minima V_min", f"{orig.get('V_min', 0):.1f} kN")
        ]
        
        for i, (param, value) in enumerate(data, start=1):
            table.cell(i, 0).text = param
            table.cell(i, 1).text = value
            
        # Stato di progetto
        doc.add_heading('6.2 Stato di Progetto', 2)
        
        mod = self.results.get('modified', {})
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.cell(0, 0).text = "Parametro"
        table.cell(0, 1).text = "Valore"
        
        data = [
            ("Rigidezza laterale K", f"{mod.get('K', 0):.1f} kN/m"),
            ("Contributo cerchiature", f"{mod.get('K_cerchiature', 0):.1f} kN/m"),
            ("Resistenza taglio V_t1", f"{mod.get('V_t1', 0):.1f} kN"),
            ("Resistenza taglio con f.f. V_t2", f"{mod.get('V_t2', 0):.1f} kN"),
            ("Resistenza minima V_min", f"{mod.get('V_min', 0):.1f} kN")
        ]
        
        for i, (param, value) in enumerate(data, start=1):
            table.cell(i, 0).text = param
            table.cell(i, 1).text = value
            
    def _add_word_verifications(self, doc):
        """Aggiunge verifiche normative Word"""
        verif = self.results.get('verification', {})
        
        # Verifica variazione rigidezza
        doc.add_heading('7.1 Verifica Variazione di Rigidezza', 2)
        
        p = doc.add_paragraph()
        stiff_var = verif.get('stiffness_variation', 0)
        stiff_ok = verif.get('stiffness_ok', False)
        
        p.add_run(f"Variazione di rigidezza: ΔK = {stiff_var:.1f}%")
        p.add_run(f"\nLimite normativo: ±15%")
        p.add_run(f"\nVerifica: ")
        
        if stiff_ok:
            p.add_run("SODDISFATTA ✓").bold = True
        else:
            p.add_run("NON SODDISFATTA ✗").bold = True
            
        # Verifica variazione resistenza
        doc.add_heading('7.2 Verifica Variazione di Resistenza', 2)
        
        p = doc.add_paragraph()
        res_var = verif.get('resistance_variation', 0)
        res_ok = verif.get('resistance_ok', False)
        
        p.add_run(f"Riduzione di resistenza: ΔV = {res_var:.1f}%")
        p.add_run(f"\nLimite normativo: max 20%")
        p.add_run(f"\nVerifica: ")
        
        if res_ok:
            p.add_run("SODDISFATTA ✓").bold = True
        else:
            p.add_run("NON SODDISFATTA ✗").bold = True
            
    def _add_word_conclusions(self, doc):
        """Aggiunge conclusioni Word"""
        verif = self.results.get('verification', {})
        is_local = verif.get('is_local', False)
        
        p = doc.add_paragraph()
        p.add_run("Sulla base delle verifiche effettuate secondo NTC 2018 § 8.4.1, ")
        
        if is_local:
            p.add_run("l'intervento è classificabile come ").bold = False
            p.add_run("INTERVENTO LOCALE").bold = True
            p.add_run(" in quanto le variazioni di rigidezza e resistenza ")
            p.add_run("rimangono entro i limiti normativi.")
            
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Pertanto ").bold = False
            p.add_run("NON").bold = True
            p.add_run(" è necessaria la valutazione della sicurezza globale della costruzione.")
        else:
            p.add_run("l'intervento ").bold = False
            p.add_run("NON").bold = True
            p.add_run(" è classificabile come intervento locale.")
            
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Si rendono necessarie ulteriori verifiche o modifiche progettuali.")
            
    def _add_word_signature(self, doc):
        """Aggiunge firma Word"""
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Luogo e data
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"Palermo, {datetime.now().strftime('%d/%m/%Y')}")
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Firma
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run("Il Progettista")
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(self.engineer_data['nome']).bold = True
        
    # Metodi PDF (implementazione base)
    def _add_pdf_header(self, story, styles):
        """Aggiunge intestazione PDF"""
        data = [
            [self.engineer_data['nome'], f"Progetto: {self.project_data.get('info', {}).get('name', 'N.D.')}"],
            [self.engineer_data['iscrizione'], f"Cliente: {self.project_data.get('info', {}).get('client', 'N.D.')}"],
            [self.engineer_data['sede_op'], f"Ubicazione: {self.project_data.get('info', {}).get('location', 'N.D.')}"]
        ]
        
        t = Table(data, colWidths=[10*cm, 8*cm])
        t.setStyle(TableStyle([
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
        ]))
        story.append(t)
        story.append(Spacer(1, 1*cm))
        
    def _add_pdf_normative(self, story, styles):
        """Aggiunge riferimenti normativi PDF"""
        story.append(Paragraph("1. RIFERIMENTI NORMATIVI", styles['Heading1']))
        story.append(Spacer(1, 0.5*cm))
        
        normative = [
            "• D.M. 17/01/2018 - Norme Tecniche per le Costruzioni (NTC 2018)",
            "• Circolare 21/01/2019 n. 7",
            "• § 8.4.1 - Interventi locali",
            "• § 8.7.1 - Costruzioni in muratura"
        ]
        
        for norm in normative:
            story.append(Paragraph(norm, styles['Normal']))
            
        story.append(Spacer(1, 0.5*cm))
        
    def _add_pdf_description(self, story, styles):
        """Aggiunge descrizione PDF"""
        story.append(Paragraph("2. DESCRIZIONE DELL'INTERVENTO", styles['Heading1']))
        story.append(Spacer(1, 0.5*cm))
        
        openings = self.project_data.get('openings_module', {}).get('openings', [])
        n_openings = len(openings)
        
        text = f"L'intervento prevede la realizzazione di {n_openings} apertura/e in parete muraria portante."
        story.append(Paragraph(text, styles['Normal']))
        story.append(Spacer(1, 0.5*cm))
        
    def _add_pdf_materials(self, story, styles):
        """Aggiunge materiali PDF"""
        story.append(Paragraph("3. CARATTERISTICHE DEI MATERIALI", styles['Heading1']))
        story.append(Spacer(1, 0.5*cm))
        
        masonry = self.project_data.get('masonry', {})
        
        data = [
            ["Parametro", "Valore"],
            ["Tipologia muratura", masonry.get('type', 'N.D.')],
            ["fcm [MPa]", str(masonry.get('fcm', 0))],
            ["τ0 [MPa]", str(masonry.get('tau0', 0))],
            ["E [MPa]", str(masonry.get('E', 0))]
        ]
        
        t = Table(data, colWidths=[10*cm, 8*cm])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.5*cm))
        
    def _add_pdf_geometry(self, story, styles):
        """Aggiunge geometria PDF"""
        story.append(Paragraph("4. GEOMETRIA DELLA PARETE", styles['Heading1']))
        story.append(Spacer(1, 0.5*cm))
        
        wall = self.project_data.get('wall', {})
        
        text = f"Dimensioni parete: L={wall.get('length', 0)}cm x h={wall.get('height', 0)}cm x t={wall.get('thickness', 0)}cm"
        story.append(Paragraph(text, styles['Normal']))
        story.append(Spacer(1, 0.5*cm))
        
    def _add_pdf_loads(self, story, styles):
        """Aggiunge carichi PDF"""
        story.append(Paragraph("5. ANALISI DEI CARICHI", styles['Heading1']))
        story.append(Spacer(1, 0.5*cm))
        
        loads = self.project_data.get('loads', {})
        
        text = f"Carico verticale: N = {loads.get('vertical', 0)} kN"
        story.append(Paragraph(text, styles['Normal']))
        story.append(Spacer(1, 0.5*cm))
        
    def _add_pdf_results(self, story, styles):
        """Aggiunge risultati PDF"""
        story.append(Paragraph("6. RISULTATI DEL CALCOLO", styles['Heading1']))
        story.append(Spacer(1, 0.5*cm))
        
        if self.results:
            orig = self.results.get('original', {})
            text = f"Rigidezza stato di fatto: K = {orig.get('K', 0):.1f} kN/m"
            story.append(Paragraph(text, styles['Normal']))
            
            text = f"Resistenza minima: V_min = {orig.get('V_min', 0):.1f} kN"
            story.append(Paragraph(text, styles['Normal']))
            
        story.append(Spacer(1, 0.5*cm))
        
    def _add_pdf_verifications(self, story, styles):
        """Aggiunge verifiche PDF"""
        story.append(Paragraph("7. VERIFICHE NORMATIVE", styles['Heading1']))
        story.append(Spacer(1, 0.5*cm))
        
        verif = self.results.get('verification', {})
        
        text = f"Variazione rigidezza: {verif.get('stiffness_variation', 0):.1f}% (limite: ±15%)"
        story.append(Paragraph(text, styles['Normal']))
        
        text = f"Variazione resistenza: {verif.get('resistance_variation', 0):.1f}% (limite: 20%)"
        story.append(Paragraph(text, styles['Normal']))
        
        story.append(Spacer(1, 0.5*cm))
        
    def _add_pdf_conclusions(self, story, styles):
        """Aggiunge conclusioni PDF"""
        story.append(Paragraph("8. CONCLUSIONI", styles['Heading1']))
        story.append(Spacer(1, 0.5*cm))
        
        verif = self.results.get('verification', {})
        if verif.get('is_local', False):
            text = "L'intervento è classificabile come INTERVENTO LOCALE secondo NTC 2018."
        else:
            text = "L'intervento NON è classificabile come intervento locale."
            
        story.append(Paragraph(text, styles['Normal']))
        story.append(Spacer(1, 0.5*cm))
        
    def _add_pdf_signature(self, story, styles):
        """Aggiunge firma PDF"""
        story.append(Spacer(1, 2*cm))
        
        text = f"Palermo, {datetime.now().strftime('%d/%m/%Y')}"
        story.append(Paragraph(text, styles['Right']))
        
        story.append(Spacer(1, 1*cm))
        
        story.append(Paragraph("Il Progettista", styles['Right']))
        story.append(Paragraph(self.engineer_data['nome'], styles['Right']))


class ReportModule(QWidget):
    """Modulo interfaccia per generazione relazioni"""
    
    # Segnale per notificare modifiche
    data_changed = pyqtSignal()
    
    def __init__(self):
        super().__init__()
        self.project_data = {}
        self.results = {}
        self.generator = ReportGenerator()
        self.setup_ui()
        
    def setup_ui(self):
        """Costruisce interfaccia"""
        layout = QVBoxLayout()
        
        # Gruppo opzioni
        options_group = QGroupBox("Opzioni Relazione")
        options_layout = QVBoxLayout()
        
        # Formato output
        format_layout = QHBoxLayout()
        format_layout.addWidget(QLabel("Formato:"))
        
        self.word_radio = QRadioButton("Word (.docx)")
        self.word_radio.setChecked(DOCX_AVAILABLE)
        self.word_radio.setEnabled(DOCX_AVAILABLE)
        format_layout.addWidget(self.word_radio)
        
        self.pdf_radio = QRadioButton("PDF")
        self.pdf_radio.setChecked(REPORTLAB_AVAILABLE and not DOCX_AVAILABLE)
        self.pdf_radio.setEnabled(REPORTLAB_AVAILABLE)
        format_layout.addWidget(self.pdf_radio)
        
        format_layout.addStretch()
        options_layout.addLayout(format_layout)
        
        # Verifica disponibilità moduli
        if not DOCX_AVAILABLE and not REPORTLAB_AVAILABLE:
            warning_label = QLabel(
                "⚠️ Installare python-docx o reportlab per generare relazioni:\n"
                "pip install python-docx reportlab"
            )
            warning_label.setStyleSheet("color: red; padding: 10px;")
            options_layout.addWidget(warning_label)
        
        # Contenuti da includere
        self.include_graphs = QCheckBox("Includi grafici")
        self.include_graphs.setChecked(True)
        options_layout.addWidget(self.include_graphs)
        
        self.include_drawings = QCheckBox("Includi disegni tecnici")
        self.include_drawings.setChecked(True)
        options_layout.addWidget(self.include_drawings)
        
        self.include_photos = QCheckBox("Includi documentazione fotografica")
        options_layout.addWidget(self.include_photos)
        
        options_group.setLayout(options_layout)
        layout.addWidget(options_group)
        
        # Anteprima
        preview_group = QGroupBox("Anteprima")
        preview_layout = QVBoxLayout()
        
        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)
        preview_layout.addWidget(self.preview_text)
        
        preview_group.setLayout(preview_layout)
        layout.addWidget(preview_group)
        
        # Pulsanti
        button_layout = QHBoxLayout()
        
        self.preview_btn = QPushButton("Aggiorna Anteprima")
        self.preview_btn.clicked.connect(self.update_preview)
        button_layout.addWidget(self.preview_btn)
        
        self.generate_btn = QPushButton("Genera Relazione")
        self.generate_btn.clicked.connect(self.generate_report)
        self.generate_btn.setEnabled(DOCX_AVAILABLE or REPORTLAB_AVAILABLE)
        button_layout.addWidget(self.generate_btn)
        
        layout.addLayout(button_layout)
        
        self.setLayout(layout)
        
    def set_data(self, project_data: Dict, results: Dict):
        """Imposta dati progetto e risultati"""
        self.project_data = project_data
        self.results = results
        self.generator.set_data(project_data, results)
        self.update_preview()
        
    def set_project_data(self, project_data: Dict):
        """Imposta solo i dati del progetto"""
        self.project_data = project_data
        self.generator.project_data = project_data
        
    def set_results(self, results: Dict):
        """Imposta solo i risultati"""
        self.results = results
        self.generator.results = results
        self.update_preview()
        
    def update_preview(self):
        """Aggiorna anteprima testuale"""
        preview = []
        
        # Intestazione
        preview.append("RELAZIONE TECNICA")
        preview.append("VERIFICA LOCALE APERTURA IN PARETE MURARIA")
        preview.append("")
        
        # Info progetto
        info = self.project_data.get('info', {})
        preview.append(f"Progetto: {info.get('name', 'N.D.')}")
        preview.append(f"Committente: {info.get('client', 'N.D.')}")
        preview.append(f"Ubicazione: {info.get('location', 'N.D.')}")
        preview.append("")
        
        # Risultati principali
        if self.results:
            verif = self.results.get('verification', {})
            preview.append("ESITO VERIFICHE:")
            
            if verif.get('is_local', False):
                preview.append("✓ INTERVENTO LOCALE VERIFICATO")
            else:
                preview.append("✗ INTERVENTO NON LOCALE")
                
            preview.append(f"Variazione rigidezza: {verif.get('stiffness_variation', 0):.1f}%")
            preview.append(f"Variazione resistenza: {verif.get('resistance_variation', 0):.1f}%")
        else:
            preview.append("(Eseguire il calcolo per visualizzare i risultati)")
            
        self.preview_text.setPlainText('\n'.join(preview))
        
    def generate_report(self):
        """Genera relazione nel formato selezionato"""
        if not self.project_data:
            QMessageBox.warning(self, "Attenzione", 
                              "Dati progetto mancanti")
            return
            
        if not self.results:
            QMessageBox.warning(self, "Attenzione", 
                              "Eseguire prima il calcolo")
            return
            
        # Determina formato
        if self.word_radio.isChecked() and DOCX_AVAILABLE:
            ext = "docx"
            filter_str = "Documenti Word (*.docx)"
        elif self.pdf_radio.isChecked() and REPORTLAB_AVAILABLE:
            ext = "pdf"
            filter_str = "Documenti PDF (*.pdf)"
        else:
            QMessageBox.warning(self, "Attenzione", 
                              "Nessun formato disponibile")
            return
            
        # Dialog salvataggio
        filename, _ = QFileDialog.getSaveFileName(
            self, "Salva Relazione",
            f"Relazione_Cerchiature_{datetime.now().strftime('%Y%m%d')}.{ext}",
            filter_str
        )
        
        if filename:
            try:
                if self.word_radio.isChecked():
                    self.generator.generate_word_report(filename)
                else:
                    self.generator.generate_pdf_report(filename)
                    
                QMessageBox.information(self, "Completato",
                                      f"Relazione generata:\n{filename}")
                                      
                # Apri file
                if os.path.exists(filename):
                    os.startfile(filename)
                    
            except Exception as e:
                QMessageBox.critical(self, "Errore",
                                   f"Errore generazione relazione:\n{str(e)}")
                                   
    def reset(self):
        """Reset del modulo"""
        self.project_data = {}
        self.results = {}
        self.generator.set_data({}, {})
        self.preview_text.clear()
        
    def collect_data(self):
        """Raccoglie dati del modulo (per compatibilità)"""
        return {
            'format': 'word' if self.word_radio.isChecked() else 'pdf',
            'include_graphs': self.include_graphs.isChecked(),
            'include_drawings': self.include_drawings.isChecked(),
            'include_photos': self.include_photos.isChecked()
        }
        
    def load_data(self, data):
        """Carica dati nel modulo (per compatibilità)"""
        if 'format' in data:
            if data['format'] == 'word' and DOCX_AVAILABLE:
                self.word_radio.setChecked(True)
            elif data['format'] == 'pdf' and REPORTLAB_AVAILABLE:
                self.pdf_radio.setChecked(True)
                
        if 'include_graphs' in data:
            self.include_graphs.setChecked(data['include_graphs'])
        if 'include_drawings' in data:
            self.include_drawings.setChecked(data['include_drawings'])
        if 'include_photos' in data:
            self.include_photos.setChecked(data['include_photos'])